from player import *
from typing import List, TextIO
from docx import *
from docx.shared import *
import datetime

SHIMPAN = []
PATH = 'C:/Users/franc/Documents/Kendo/{} UofT Tournament/'.format(datetime.datetime.now().year)
DATE = 'Saturday, March 14th'
PRESIDENT = 'Luke Pham'


def find_replace(document, find, replace):
    for p in document.paragraphs:
        p.text = p.text.replace(find, replace)


def get_shimpan_list(player_list: List[Player]) -> List[Player]:
    shimpan_list = []

    for player in player_list:
        if player.shimpan:
            shimpan_list.append(player)

    return shimpan_list


def duplicate_template(shimpan_list: List[Player]):

    for shimpan in shimpan_list:
        if shimpan.rank >= 6:
            shutil.copy('Shimpan Letters/5Dan+ Shimpan Parking Letter.docx',
                         PATH + 'Shimpan/Print/{} {}.docx'.format(shimpan.club, shimpan.name))
            personalize_letter(PATH + 'Shimpan/Print/{} {}.docx'.format(shimpan.club, shimpan.name), shimpan, CLUB_LIST)
        elif shimpan.rank == 5:
            shutil.copy('Shimpan Letters/5Dan+ Shimpan Parking Letter.docx',
                         PATH + 'Shimpan/Email/{} {}.docx'.format(shimpan.club, shimpan.name))
            personalize_letter(PATH + 'Shimpan/Email/{} {}.docx'.format(shimpan.club, shimpan.name), shimpan, CLUB_LIST)
        else:
            shutil.copy('Shimpan Letters/4Dan Shimpan No Parking Letter.docx',
                         PATH + 'Shimpan/Email/{} {}.docx'.format(shimpan.club, shimpan.name))
            personalize_letter(PATH + 'Shimpan/Email/{} {}.docx'.format(shimpan.club, shimpan.name), shimpan, CLUB_LIST)


def personalize_letter (document, shimpan, club_list):

    doc = Document(document)

    #for i, line in enumerate(doc.paragraphs):
    #    print (i, line.text)

    doc.paragraphs[0].text = doc.paragraphs[0].text.replace('(sensei)', shimpan.name)
    doc.paragraphs[1].text = doc.paragraphs[1].text.replace('(club)', club_list[shimpan.club].fullname)
    doc.paragraphs[3].text = doc.paragraphs[3].text.replace('(sensei)', shimpan.name)
    doc.paragraphs[6].text = doc.paragraphs[6].text.replace('(date)', DATE)
    doc.paragraphs[6].text = doc.paragraphs[6].text.replace('(year)', str(datetime.datetime.now().year))

    try:
        doc.paragraphs[16].text = doc.paragraphs[16].text.replace('(president)', PRESIDENT)
    except:
        doc.paragraphs[13].text = doc.paragraphs[13].text.replace('(president)', PRESIDENT)
    

    doc.save(document)


if __name__ == '__main__':    
    CLUB_LIST = {}
    players = []

    a = convert_to_Player('UOT.csv')
    b = convert_to_Player('JCC.csv')
    players.extend(a)
    players.extend(b)

    CLUB_LIST['UOT'] = Dojo('UOT', read_club_file_info('UOT.csv'), a)
    CLUB_LIST['JCC'] = Dojo('JCC', read_club_file_info('JCC.csv'), b)

    SHIMPAN = get_shimpan_list(players)
    duplicate_template(SHIMPAN)